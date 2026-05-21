"""
Document ingestion worker — format-aware chunking.

Pipeline (non-markdown)
-----------------------
1. ``pypdf`` or ``python-docx`` extracts text from PDF / DOCX / TXT.
2. All extracted text is joined and split with a single
   ``RecursiveCharacterTextSplitter`` using a configurable chunk size and
   overlap (env vars ``CHUNK_SIZE`` / ``CHUNK_OVERLAP``).

Pipeline (markdown .md)
-----------------------
1. The file is read directly as UTF-8 text.
2. A header-aware parser splits the document on H1–H6 boundaries, building
   a tree of sections.  Each section becomes one chunk; sections larger than
   ``MARKDOWN_CHUNK_SIZE`` are sub-split with the same fallback splitter.

3. Chunks are upserted into **Qdrant** (vector search) and persisted to a
   MongoDB ``document_chunks`` collection for auditability.
4. The uploaded file is persisted to the **file-service** (S3).
5. A document catalog record is saved to ``ingested_documents``.
"""

import logging
import os
import re
import unicodedata
from pathlib import Path
from typing import Any, Dict, List, Tuple

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config.settings import settings
from app.repositories.chunk_repository import get_chunk_repository
from app.repositories.document_repository import get_document_repository
from app.services.file_service_client import get_file_service_client
from app.services.task_manager import TaskStatus, get_task_manager
from app.services.vector_db import get_vector_service
from app.utils.markdown_chunker import build_chunks_from_markdown

logger = logging.getLogger(__name__)

# ── Single shared splitter ────────────────────────────────────────────────────

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=getattr(settings, "CHUNK_SIZE", 1000),
    chunk_overlap=getattr(settings, "CHUNK_OVERLAP", 200),
    separators=["\n\n", "\n", ". ", " ", ""],
    length_function=len,
)

# ── Text cleaning ─────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Normalise whitespace and strip non-printable characters.

    Vietnamese characters (including combining diacritics) are preserved by
    first normalising to NFC form, then filtering only truly non-printable
    control characters (categories Cc, Cf, Cs, Co, Cn).
    """
    text = unicodedata.normalize("NFC", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n+", "\n", text)
    text = "".join(
        ch for ch in text
        if unicodedata.category(ch)[0] != "C" or ch in "\n\t"
    )
    return text.strip()


# ── Format-specific text extraction ───────────────────────────────────────────

def _extract_text_blocks(file_path: Path) -> List[str]:
    """Extract text blocks from a document using lightweight parsers.

    Returns a list of non-empty cleaned text strings — one per element.
    """
    suffix = file_path.suffix.lower()
    blocks: List[str] = []
    
    try:
        if suffix == ".pdf":
            blocks = _extract_pdf(str(file_path))
        elif suffix in [".docx", ".doc"]:
            blocks = _extract_docx(str(file_path))
        elif suffix == ".txt":
            blocks = _extract_txt(str(file_path))
        else:
            logger.warning("Unsupported file type: %s", suffix)
            return []
            
        logger.info("Extracted %d text blocks from %s", len(blocks), file_path.name)
    except Exception as e:
        logger.error("Failed to extract text from %s: %s", file_path.name, e)
        raise
    
    return [b for b in blocks if b.strip()]


def _extract_pdf(file_path: str) -> List[str]:
    """Extract text from PDF using pypdf."""
    from pypdf import PdfReader
    
    reader = PdfReader(file_path)
    blocks = []
    
    for page_num, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text:
                text = _clean(text)
                if text:
                    blocks.append(text)
        except Exception as e:
            logger.warning(f"Failed to extract page {page_num}: {e}")
            continue
    
    return blocks


def _extract_docx(file_path: str) -> List[str]:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    
    doc = Document(file_path)
    blocks = []
    current_para = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current_para.append(text)
        elif current_para:
            # End of a paragraph group
            combined = " ".join(current_para)
            if combined.strip():
                blocks.append(_clean(combined))
            current_para = []
    
    # Don't forget last paragraph
    if current_para:
        combined = " ".join(current_para)
        if combined.strip():
            blocks.append(_clean(combined))
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(_clean(row_text))
    
    return blocks


def _extract_txt(file_path: str) -> List[str]:
    """Extract text from plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    
    if not content:
        return []
    
    # Split by double newlines (paragraphs)
    paragraphs = content.split("\n\n")
    return [_clean(p.strip()) for p in paragraphs if p.strip()]


# ── Chunking ──────────────────────────────────────────────────────────────────

def _build_chunks(
    blocks: List[str],
    metadata: Dict[str, Any],
    file_hash: str,
    source_name: str,
) -> List[Dict[str, Any]]:
    """Split text blocks into fixed-size chunks and attach metadata."""
    full_text = "\n\n".join(blocks)
    splits = _splitter.split_text(full_text)

    base_meta: Dict[str, Any] = {
        **metadata,
        "file_hash": file_hash,
        "source": source_name,
    }
    return [{"text": s, "metadata": base_meta.copy()} for s in splits]


# ── Preview-only helper (no persistence) ─────────────────────────────────────

def preview_document(file_path: Path) -> Tuple[List[Dict[str, Any]], List[str]]:
    """Parse and chunk a document without persisting anything.

    For .md files the section list is populated with full heading paths.
    For other formats the section list is empty.
    """
    is_markdown = file_path.suffix.lower() == ".md"

    if is_markdown:
        from app.utils.markdown_chunker import preview_markdown
        return preview_markdown(file_path)

    blocks = _extract_text_blocks(file_path)
    if not blocks:
        return [], []
    chunks = _build_chunks(blocks, {}, "", file_path.name)
    return chunks, []


# ── Main entry point (called as a BackgroundTask) ────────────────────────────

async def process_document(
    file_path: Path,
    metadata: Dict[str, Any],
    file_hash: str,
    task_id: str,
) -> None:
    """Background task that ingests a document end-to-end.

    Steps (non-markdown)
    --------------------
    1. Text extraction  (pypdf / python-docx)
    2. Fixed-size chunking
    3. Persist chunks to MongoDB (``document_chunks``)
    4. Embed & upsert into Qdrant
    5. Upload file to file-service (S3)
    6. Save document catalog record (``ingested_documents``)
    7. Clean up temp file

    Steps (markdown .md)
    ---------------------
    1. Parse by heading structure (semantic sections)
    2. Sub-split oversized sections
    3. Persist chunks to MongoDB
    4. Embed & upsert into Qdrant
    5. Upload file to file-service (S3)
    6. Save document catalog record
    7. Clean up temp file
    """
    task_manager = get_task_manager()
    task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.start")

    is_markdown = file_path.suffix.lower() == ".md"

    try:
        if is_markdown:
            # ── Markdown: header-based chunking ──────────────────────────────
            task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.section_chunk")
            chunks = build_chunks_from_markdown(file_path, metadata, file_hash)
            source_name = file_path.name
        else:
            # ── PDF / DOCX / TXT: lightweight parser + fixed-size chunking ───
            task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.layout_parse")
            blocks = _extract_text_blocks(file_path)

            if not blocks:
                logger.warning("No text extracted from %s", file_path.name)
                task_manager.update_task(task_id, TaskStatus.FAILED, message="task.processing.empty")
                return

            task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.section_chunk")
            source_name = file_path.name
            chunks = _build_chunks(blocks, metadata, file_hash, source_name)

        if not chunks:
            logger.warning("Chunking produced 0 chunks for %s", file_path.name)
            task_manager.update_task(task_id, TaskStatus.FAILED, message="task.processing.empty")
            return

        logger.info("Produced %d chunks for %s", len(chunks), file_path.name)

        # 4. Embed & index in Qdrant ───────────────────────────────────────────
        task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.index")
        docs = [Document(page_content=c["text"], metadata=c["metadata"]) for c in chunks]
        vector_service = get_vector_service()
        add_result = vector_service.add_documents(docs)
        # Capture point IDs returned by Qdrant so each chunk can be linked
        # back to its vector record for the source-document viewer.
        # langchain-qdrant >= 0.1.0 returns a list of str point IDs;
        # older versions returned structs with a .id attribute.  Handle both.
        if add_result:
            point_ids = [rec.id if hasattr(rec, "id") else str(rec) for rec in add_result]
        else:
            point_ids = []
        logger.info("Indexed %d chunks for %s — pointIds captured: %s", len(docs), file_path.name, len(point_ids))

        # 3. Persist chunks to MongoDB ─────────────────────────────────────────
        task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.store_chunks")
        chunk_repo = get_chunk_repository()
        chunk_repo.store_chunks(
            document_id=file_hash,
            chunks=chunks,
            source_file=metadata.get("original_filename", source_name),
            point_ids=point_ids,
        )

        # 5. Upload to file-service (S3) ───────────────────────────────────────
        task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.upload_file")
        file_service_id = None
        file_service_s3_key = None
        try:
            client = get_file_service_client()
            content_type = metadata.get("content_type", "application/octet-stream")
            file_service_id, file_service_s3_key = client.upload_file(
                file_path, content_type=content_type,
            )
        except Exception as upload_exc:
            logger.warning("File-service upload failed (non-fatal): %s", upload_exc)

        # 6. Save document catalog record ──────────────────────────────────────
        task_manager.update_task(task_id, TaskStatus.PROCESSING, message="task.processing.save_catalog")
        try:
            doc_repo = get_document_repository()
            file_size = file_path.stat().st_size if file_path.exists() else None

            # Extract section titles for markdown documents
            if is_markdown:
                sections = list(dict.fromkeys(
                    c["metadata"].get("section_full_path", "")
                    for c in chunks
                    if c["metadata"].get("section_full_path")
                ))
            else:
                sections = []

            doc_repo.save_document({
                "document_id": file_hash,
                "original_filename": metadata.get("original_filename", source_name),
                "content_type": metadata.get("content_type"),
                "file_size": file_size,
                "category": metadata.get("category"),
                "variety": metadata.get("variety"),
                "user_id": metadata.get("user_id"),
                "file_service_id": file_service_id,
                "file_service_s3_key": file_service_s3_key,
                "chunk_count": len(chunks),
                "sections": sections,
                "status": "ingested",
            })
        except Exception as catalog_exc:
            logger.warning("Document catalog save failed (non-fatal): %s", catalog_exc)

        task_manager.update_task(task_id, TaskStatus.COMPLETED, message="task.completed")

    except Exception as exc:
        logger.exception("Error processing document %s: %s", file_path, exc)
        task_manager.update_task(
            task_id, TaskStatus.FAILED, error=str(exc), message="task.processing.failed",
        )
    finally:
        # 7. Cleanup temp file ─────────────────────────────────────────────────
        if file_path.exists():
            os.remove(file_path)
            logger.info("Cleaned up temp file: %s", file_path)
